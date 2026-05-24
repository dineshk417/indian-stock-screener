import io
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Design constants ────────────────────────────────────────────────────────
_FONT        = "Calibri"
_FONT_SERIF  = "Cambria"   # for name only
_BLACK       = RGBColor(0x00, 0x00, 0x00)
_DARK        = RGBColor(0x1A, 0x1A, 0x1A)
_MID         = RGBColor(0x3A, 0x3A, 0x3A)
_LIGHT       = RGBColor(0x55, 0x55, 0x55)

# Content width on A4 with 1.9cm margins each side ≈ 9780 twips
_PAGE_W_TWIPS = 9780
# Left column ~72%, right ~28%
_COL_L = Cm(12.2)
_COL_R = Cm(4.6)


# ── Low-level XML helpers ────────────────────────────────────────────────────

def _set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _no_space_before_table(table):
    """Remove the default paragraph gap that Word inserts before tables."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblStyle = OxmlElement("w:tblCellSpacing")
    tblStyle.set(qn("w:w"), "0")
    tblStyle.set(qn("w:type"), "dxa")
    tblPr.append(tblStyle)


def _paragraph_border_bottom(paragraph, color="000000", size=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"), "single")
    btm.set(qn("w:sz"), str(size))
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), color)
    pBdr.append(btm)
    pPr.append(pBdr)


def _space(p, before=0, after=0, line_pt=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line_pt:
        pf.line_spacing = Pt(line_pt)


def _run(para, text, bold=False, italic=False, size=10, color=None, font=None):
    r = para.add_run(text)
    r.font.name        = font or _FONT
    r.font.size        = Pt(size)
    r.font.bold        = bold
    r.font.italic      = italic
    r.font.color.rgb   = color or _BLACK
    return r


# ── Section header ───────────────────────────────────────────────────────────

def _section_header(doc, title):
    p = doc.add_paragraph()
    _space(p, before=9, after=2)
    _paragraph_border_bottom(p, color="000000", size=4)
    _run(p, title.upper(), bold=True, size=10.5, color=_BLACK)
    return p


# ── Entry (job / education) ──────────────────────────────────────────────────

def _add_entry(doc, role, org, duration, location, bullets):
    # Two-column borderless table: Role | Date
    tbl = doc.add_table(rows=2, cols=2)
    _remove_table_borders(tbl)
    _no_space_before_table(tbl)

    # Fix column widths
    for row in tbl.rows:
        row.cells[0].width = _COL_L
        row.cells[1].width = _COL_R
        for cell in row.cells:
            _set_cell_margins(cell, top=0, start=0, bottom=0, end=0)

    # Row 0 – Role (bold) | Duration (right-aligned)
    r0c0 = tbl.rows[0].cells[0].paragraphs[0]
    r0c1 = tbl.rows[0].cells[1].paragraphs[0]
    _run(r0c0, role, bold=True, size=10, color=_DARK)
    r0c0.paragraph_format.space_before = Pt(4)
    r0c0.paragraph_format.space_after  = Pt(0)
    if duration:
        _run(r0c1, duration, size=9.5, color=_MID)
        r0c1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0c1.paragraph_format.space_before = Pt(4)
        r0c1.paragraph_format.space_after  = Pt(0)

    # Row 1 – Org (italic) | Location (right-aligned)
    r1c0 = tbl.rows[1].cells[0].paragraphs[0]
    r1c1 = tbl.rows[1].cells[1].paragraphs[0]
    if org:
        _run(r1c0, org, italic=True, size=9.5, color=_MID)
        r1c0.paragraph_format.space_before = Pt(0)
        r1c0.paragraph_format.space_after  = Pt(1)
    if location:
        _run(r1c1, location, size=9, color=_LIGHT)
        r1c1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1c1.paragraph_format.space_before = Pt(0)
        r1c1.paragraph_format.space_after  = Pt(1)

    # Bullet points
    for bullet in bullets:
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent        = Cm(0.45)
        bp.paragraph_format.first_line_indent  = Pt(-10)
        bp.paragraph_format.space_before       = Pt(1)
        bp.paragraph_format.space_after        = Pt(1)
        _run(bp, "• " + bullet, size=9.5, color=_DARK)


# ── Main builder ─────────────────────────────────────────────────────────────

def build_resume_docx(structure: dict) -> bytes:
    doc = Document()

    # Page setup
    for sec in doc.sections:
        sec.top_margin    = Cm(1.9)
        sec.bottom_margin = Cm(1.9)
        sec.left_margin   = Cm(1.9)
        sec.right_margin  = Cm(1.9)

    # Reset Normal style spacing
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(0)
    normal.font.name = _FONT
    normal.font.size = Pt(10)

    # ── Name ──
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _space(name_p, before=0, after=3)
    name_text = structure.get("name") or "Your Name"
    _run(name_p, name_text.upper(), bold=True, size=18, color=_BLACK, font=_FONT_SERIF)

    # ── Contact bar ──
    contacts = [c for c in structure.get("contact", []) if c and c.strip()]
    if contacts:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _space(cp, before=0, after=6)
        _run(cp, "  •  ".join(contacts), size=9, color=_MID)

    # ── Summary ──
    summary = (structure.get("summary") or "").strip()
    if summary:
        _section_header(doc, "Summary")
        sp = doc.add_paragraph()
        _space(sp, before=3, after=4)
        _run(sp, summary, size=9.5, color=_DARK)

    # ── Sections ──
    for section in structure.get("sections", []):
        title = section.get("title", "")
        kind  = section.get("type", "text")
        _section_header(doc, title)

        if kind == "entries":
            for e in section.get("entries", []):
                _add_entry(
                    doc,
                    role     = e.get("role", ""),
                    org      = e.get("org", ""),
                    duration = e.get("duration", ""),
                    location = e.get("location", ""),
                    bullets  = e.get("bullets", []),
                )

        elif kind == "skills":
            for cat in section.get("categories", []):
                kp = doc.add_paragraph()
                _space(kp, before=2, after=2)
                label = (cat.get("label") or "").strip()
                items = (cat.get("items") or "").strip()
                if label:
                    _run(kp, label + ":  ", bold=True, size=9.5, color=_DARK)
                _run(kp, items, size=9.5, color=_DARK)

        else:
            tp = doc.add_paragraph()
            _space(tp, before=2, after=4)
            _run(tp, section.get("content", ""), size=9.5, color=_DARK)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Cover letter (plain, professional) ───────────────────────────────────────

def text_to_docx_bytes(text: str, title: str = "") -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)
    if title:
        h = doc.add_heading(title, level=1)
        h.runs[0].font.name  = _FONT_SERIF
        h.runs[0].font.size  = Pt(14)
        h.runs[0].font.color.rgb = _BLACK
        h.paragraph_format.space_after = Pt(12)
    for line in text.splitlines():
        p = doc.add_paragraph(line or " ")
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.name = _FONT
            r.font.size = Pt(10.5)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
